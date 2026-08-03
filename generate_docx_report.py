import os
import docx
import math
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

# ==================== PIL DRAWING UTILITIES (HIGH-DPI SCALED BY 2X) ====================

def get_text_size(draw, text, font):
    try:
        bbox = draw.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    except AttributeError:
        try:
            return draw.textsize(text, font=font)
        except Exception:
            return len(text) * 16, 26

def load_font(size, bold=False):
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    try:
        return ImageFont.truetype(font_name, size)
    except Exception:
        try:
            return ImageFont.truetype("calibri.ttf", size)
        except Exception:
            return ImageFont.load_default()

def draw_centered_text(draw, text, x1, y1, x2, y2, font, fill="black"):
    w = x2 - x1
    h = y2 - y1
    lines = text.split('\n')
    line_heights = []
    line_widths = []
    for line in lines:
        lw, lh = get_text_size(draw, line, font)
        line_widths.append(lw)
        line_heights.append(lh)
    total_height = sum(line_heights) + (len(lines) - 1) * 8
    cy = y1 + (h - total_height) // 2
    for idx, line in enumerate(lines):
        lw = line_widths[idx]
        lh = line_heights[idx]
        cx = x1 + (w - lw) // 2
        draw.text((cx, cy), line, font=font, fill=fill)
        cy += lh + 8

def draw_database(draw, x1, y1, x2, y2, outline="black", fill="white"):
    h = y2 - y1
    w = x2 - x1
    eh = int(h * 0.18)
    draw.rectangle([x1, y1 + eh//2, x2, y2 - eh//2], fill=fill, outline=outline, width=4)
    draw.ellipse([x1, y1, x2, y1 + eh], fill=fill, outline=outline, width=4)
    draw.ellipse([x1, y2 - eh, x2, y2], fill=fill, outline=outline, width=4)

def draw_circle(draw, cx, cy, r, outline="black", fill="white"):
    draw.ellipse([cx - r, cy - r, cx + r, cy + r], fill=fill, outline=outline, width=4)

def draw_diamond(draw, x1, y1, x2, y2, outline="black", fill="white"):
    cx = (x1 + x2) // 2
    cy = (y1 + y2) // 2
    points = [(cx, y1), (x2, cy), (cx, y2), (x1, cy)]
    draw.polygon(points, fill=fill, outline=outline)
    draw.line([points[0], points[1], points[2], points[3], points[0]], fill=outline, width=4)

def draw_arrow(draw, x1, y1, x2, y2, outline="black"):
    draw.line([x1, y1, x2, y2], fill=outline, width=5)
    angle = math.atan2(y2 - y1, x2 - x1)
    arrow_len = 28
    px1 = x2 - arrow_len * math.cos(angle - math.pi/6)
    py1 = y2 - arrow_len * math.sin(angle - math.pi/6)
    px2 = x2 - arrow_len * math.cos(angle + math.pi/6)
    py2 = y2 - arrow_len * math.sin(angle + math.pi/6)
    draw.polygon([(x2, y2), (px1, py1), (px2, py2)], fill=outline)

def draw_arrow_with_label_offset(draw, x1, y1, x2, y2, label, font, x_offset=0, y_offset=0, outline="black", bg_fill="white"):
    draw_arrow(draw, x1, y1, x2, y2, outline=outline)
    mx = (x1 + x2) // 2 + x_offset
    my = (y1 + y2) // 2 + y_offset
    lines = label.split('\n')
    line_widths = []
    line_heights = []
    for line in lines:
        lw, lh = get_text_size(draw, line, font)
        line_widths.append(lw)
        line_heights.append(lh)
    max_w = max(line_widths)
    total_h = sum(line_heights) + (len(lines) - 1) * 8
    rx1 = mx - max_w // 2 - 8
    ry1 = my - total_h // 2 - 4
    rx2 = mx + max_w // 2 + 8
    ry2 = my + total_h // 2 + 4
    draw.rectangle([rx1, ry1, rx2, ry2], fill=bg_fill)
    draw_centered_text(draw, label, rx1, ry1, rx2, ry2, font, fill=outline)

# ==================== HIGH-DPI DIAGRAM GENERATORS ====================

def generate_architecture_diagram():
    img = Image.new("RGB", (1600, 1040), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(30)
    
    # Present Tier
    draw.rectangle([100, 60, 1500, 260], fill="#F9FAFB", outline="black", width=4)
    draw_centered_text(draw, "PRESENTATION TIER (React SPA / Vanilla CSS)\n- Dashboards (Admin, Agent, Officer roles)\n- Dynamic Hindi/English translation, Accessibility Scaling (A-, A, A+)\n- Client-Side jsPDF Document Generation", 100, 60, 1500, 260, f_body, fill="black")
    
    # Offline Fallback
    draw.rectangle([1000, 320, 1500, 440], fill="#F3F4F6", outline="black", width=4)
    draw_centered_text(draw, "OFFLINE LAYER (LocalStorage)\n- Axios Mock Adapter intercepts API requests", 1000, 320, 1500, 440, f_body, fill="black")
    
    # Middleware
    draw.rectangle([100, 500, 1500, 700], fill="#F9FAFB", outline="black", width=4)
    draw_centered_text(draw, "APPLICATION MIDDLEWARE TIER (Node.js & Express.js REST APIs)\n- Security: Bcrypt Password Hashing, JWT session verification\n- Controllers: Vessel Registry CRUD, Workflow Stepper State engine\n- Security Audit Logs recorder", 100, 500, 1500, 700, f_body, fill="black")
    
    # Data Tier
    draw.rectangle([100, 820, 1500, 980], fill="#F3F4F6", outline="black", width=4)
    draw_centered_text(draw, "DATA STORAGE TIER (MongoDB Atlas Cloud Cluster)\n- Users Collection   - Vessels Collection\n- Journeys Collection   - Audit Trails Logs", 100, 820, 1500, 980, f_body, fill="black")
    
    # Arrows
    draw_arrow(draw, 500, 260, 500, 500, outline="black")
    draw_arrow(draw, 800, 260, 1000, 380, outline="black")
    draw_arrow(draw, 1000, 420, 800, 500, outline="black")
    draw_arrow(draw, 800, 700, 800, 820, outline="black")
    
    img.save("architecture.png")

def generate_cfd():
    img = Image.new("RGB", (1400, 800), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(22)
    f_title = load_font(24, bold=True)
    
    draw_circle(draw, 700, 340, 110, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "PORT DIGITAL\nCLEARANCE\nSYSTEM PORTAL", 590, 230, 810, 450, f_title, fill="black")
    
    # Entities
    draw.rectangle([50, 280, 300, 400], fill="white", outline="black", width=4)
    draw_centered_text(draw, "SHIP AGENT", 50, 280, 300, 400, f_title, fill="black")
    
    draw.rectangle([1100, 280, 1350, 400], fill="white", outline="black", width=4)
    draw_centered_text(draw, "OFFICERS\n(PHO, Customs, Traffic)", 1100, 280, 1350, 400, f_title, fill="black")
    
    draw.rectangle([520, 620, 880, 740], fill="white", outline="black", width=4)
    draw_centered_text(draw, "ADMINISTRATOR", 520, 620, 880, 740, f_title, fill="black")
    
    draw_arrow_with_label_offset(draw, 300, 310, 590, 310, "Voyage Requests\n& Vessel Data", f_body, y_offset=-45)
    draw_arrow_with_label_offset(draw, 590, 370, 300, 370, "PDF Certificates", f_body, y_offset=45)
    
    draw_arrow_with_label_offset(draw, 810, 310, 1100, 310, "Voyages Pending\nReview", f_body, y_offset=-45)
    draw_arrow_with_label_offset(draw, 1100, 370, 810, 370, "Clearance\nDecisions", f_body, y_offset=45)
    
    draw_arrow_with_label_offset(draw, 650, 450, 650, 620, "Audit Logs", f_body, x_offset=-120)
    draw_arrow_with_label_offset(draw, 750, 620, 750, 450, "Approvals", f_body, x_offset=120)
    
    img.save("cfd.png")

def generate_dfd_l0():
    img = Image.new("RGB", (1300, 400), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(22)
    f_title = load_font(24, bold=True)
    
    draw.rectangle([80, 140, 280, 260], fill="white", outline="black", width=4)
    draw_centered_text(draw, "USER", 80, 140, 280, 260, f_title, fill="black")
    
    draw_circle(draw, 600, 200, 100, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Auth\nProcess", 500, 100, 700, 300, f_title, fill="black")
    
    draw_database(draw, 1020, 140, 1220, 260, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Users DB", 1020, 170, 1220, 250, f_title, fill="black")
    
    draw_arrow_with_label_offset(draw, 280, 170, 500, 170, "Login Credentials", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 500, 230, 280, 230, "JWT Token", f_body, y_offset=35)
    draw_arrow_with_label_offset(draw, 700, 200, 1020, 200, "Query / Verify", f_body, y_offset=-35)
    
    img.save("dfd_l0.png")

def generate_dfd_l1():
    img = Image.new("RGB", (1700, 1100), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(22)
    f_title = load_font(24, bold=True)
    
    # Column 1: Entities
    draw.rectangle([100, 250, 350, 370], fill="white", outline="black", width=4)
    draw_centered_text(draw, "SHIP AGENT", 100, 250, 350, 370, f_title, fill="black")
    
    draw.rectangle([100, 700, 350, 820], fill="white", outline="black", width=4)
    draw_centered_text(draw, "OFFICERS", 100, 700, 350, 820, f_title, fill="black")
    
    # Column 2: Processes
    draw_circle(draw, 600, 160, 90, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Vessel\nCRUD", 510, 70, 690, 250, f_title, fill="black")
    
    draw_circle(draw, 600, 460, 90, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Voyage\nFiling", 510, 370, 690, 550, f_title, fill="black")
    
    draw_circle(draw, 600, 760, 90, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Clearance\nStepper", 510, 670, 690, 850, f_title, fill="black")
    
    # Column 3: Data Stores
    draw_database(draw, 950, 100, 1110, 220, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Vessels\nStore", 950, 130, 1110, 210, f_title, fill="black")
    
    draw_database(draw, 950, 400, 1110, 520, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Voyages\nStore", 950, 430, 1110, 510, f_title, fill="black")
    
    draw_circle(draw, 1030, 760, 90, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Audit\nLogger", 940, 670, 1120, 850, f_title, fill="black")
    
    # Column 4: Outputs
    draw.rectangle([1350, 160, 1600, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "PDF Reports", 1350, 160, 1600, 280, f_title, fill="black")
    
    draw_circle(draw, 1350, 460, 90, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Certificate\nGen", 1260, 370, 1440, 550, f_title, fill="black")
    
    draw_database(draw, 1350, 700, 1510, 820, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Audit\nStore", 1350, 730, 1510, 810, f_title, fill="black")
    
    draw_arrow_with_label_offset(draw, 350, 280, 510, 160, "IMO, Flag", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 350, 340, 510, 460, "Voyage Info", f_body, y_offset=35)
    draw_arrow_with_label_offset(draw, 350, 760, 510, 760, "Decision, Note", f_body, y_offset=-35)
    
    draw_arrow_with_label_offset(draw, 690, 160, 950, 160, "Save Vessel", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 690, 460, 950, 460, "Save Voyage", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 664, 696, 950, 520, "Update Status", f_body, x_offset=-30, y_offset=35)
    draw_arrow_with_label_offset(draw, 690, 760, 940, 760, "Log Event", f_body, y_offset=-35)
    
    draw_arrow_with_label_offset(draw, 1120, 760, 1350, 760, "Save Log", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 1110, 460, 1260, 460, "Fetch Voyage", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 1350, 370, 1475, 280, "Generate", f_body, y_offset=-35)
    
    img.save("dfd_l1.png")

def generate_dfd_l2():
    img = Image.new("RGB", (1500, 440), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(20)
    f_title = load_font(24, bold=True)
    
    draw_circle(draw, 160, 220, 80, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "PHO\nReview", 80, 140, 240, 300, f_title, fill="black")
    
    draw_circle(draw, 520, 220, 80, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Customs\nAudit", 440, 140, 600, 300, f_title, fill="black")
    
    draw_circle(draw, 880, 220, 80, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Traffic\nControl", 800, 140, 960, 300, f_title, fill="black")
    
    draw_circle(draw, 1240, 220, 80, outline="black", fill="#F9FAFB")
    draw_centered_text(draw, "Certificate\nGen", 1160, 140, 1320, 300, f_title, fill="black")
    
    draw_arrow_with_label_offset(draw, 240, 220, 440, 220, "Health Approved", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 600, 220, 800, 220, "Customs Approved", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 960, 220, 1160, 220, "Berth Issued", f_body, y_offset=-35)
    
    img.save("dfd_l2.png")

def generate_decomposition_diagram():
    img = Image.new("RGB", (1600, 850), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(22)
    f_title = load_font(24, bold=True)

    draw.rectangle([550, 30, 1050, 110], fill="#F3F4F6", outline="black", width=4)
    draw_centered_text(draw, "PORT DIGITAL CLEARANCE SYSTEM", 550, 30, 1050, 110, f_title, fill="black")

    # Subsystems
    draw.rectangle([150, 200, 480, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "IDENTITY & SECURITY", 150, 200, 480, 280, f_title, fill="black")
    
    draw.rectangle([650, 200, 950, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "VOYAGE CONTROL", 650, 200, 950, 280, f_title, fill="black")
    
    draw.rectangle([1120, 200, 1450, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "SYSTEM UTILITIES", 1120, 200, 1450, 280, f_title, fill="black")

    leafs_a = ["User Registration", "Admin Approval Queue", "Bcrypt & TOTP 2FA", "Inactivity Timeout (160s)"]
    leafs_b = ["Vessel Registry CRUD", "Voyage Entry Filing", "Stepper Review Pipeline", "Bilingual jsPDF Generator"]
    leafs_c = ["Public QR Verify Portal", "SVG Dashboards & Metrics", "Security Audit Logger", "Offline Mock DB Adapter"]

    for i in range(4):
        draw.rectangle([150, 340 + i*120, 480, 420 + i*120], fill="white", outline="black", width=4)
        draw_centered_text(draw, leafs_a[i], 150, 340 + i*120, 480, 420 + i*120, f_body, fill="black")
        
        draw.rectangle([650, 340 + i*120, 950, 420 + i*120], fill="white", outline="black", width=4)
        draw_centered_text(draw, leafs_b[i], 650, 340 + i*120, 950, 420 + i*120, f_body, fill="black")
        
        draw.rectangle([1120, 340 + i*120, 1450, 420 + i*120], fill="white", outline="black", width=4)
        draw_centered_text(draw, leafs_c[i], 1120, 340 + i*120, 1450, 420 + i*120, f_body, fill="black")

    # Connection Lines
    draw.line([800, 110, 800, 150], fill="black", width=4)
    draw.line([315, 150, 1285, 150], fill="black", width=4)
    draw.line([315, 150, 315, 200], fill="black", width=4)
    draw.line([800, 150, 800, 200], fill="black", width=4)
    draw.line([1285, 150, 1285, 200], fill="black", width=4)

    draw.line([315, 280, 315, 340], fill="black", width=4)
    draw.line([800, 280, 800, 340], fill="black", width=4)
    draw.line([1285, 280, 1285, 340], fill="black", width=4)

    for i in range(3):
        draw.line([315, 420 + i*120, 315, 340 + (i+1)*120], fill="black", width=4)
        draw.line([800, 420 + i*120, 800, 340 + (i+1)*120], fill="black", width=4)
        draw.line([1285, 420 + i*120, 1285, 340 + (i+1)*120], fill="black", width=4)

    img.save("decomp.png")

def generate_er_diagram():
    img = Image.new("RGB", (1700, 1100), "white")
    draw = ImageDraw.Draw(img)
    f_body = load_font(22)
    f_title = load_font(24, bold=True)
    
    # Entities
    draw.rectangle([100, 160, 380, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "USER ACCOUNT\n- _id (PK)\n- username, role", 100, 160, 380, 280, f_body, fill="black")
    
    draw.rectangle([1100, 160, 1440, 280], fill="white", outline="black", width=4)
    draw_centered_text(draw, "VESSEL REGISTRY\n- _id (PK)\n- name, imoNumber", 1100, 160, 1440, 280, f_body, fill="black")
    
    draw.rectangle([1100, 720, 1440, 880], fill="white", outline="black", width=4)
    draw_centered_text(draw, "VOYAGE (JOURNEY)\n- _id (PK)\n- clearances, status\n- ilhReceiptNo", 1100, 720, 1440, 880, f_body, fill="black")
    
    draw.rectangle([100, 720, 380, 840], fill="white", outline="black", width=4)
    draw_centered_text(draw, "AUDIT TRAILS\n- _id (PK)\n- timestamp, action", 100, 720, 380, 840, f_body, fill="black")
    
    # Relationships
    draw_diamond(draw, 620, 160, 820, 280, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Registers", 620, 160, 820, 280, f_title, fill="black")
    
    draw_diamond(draw, 620, 440, 820, 560, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Files", 620, 440, 820, 560, f_title, fill="black")
    
    draw_diamond(draw, 1170, 440, 1370, 560, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Embeds", 1170, 440, 1370, 560, f_title, fill="black")
    
    draw_diamond(draw, 140, 440, 340, 560, outline="black", fill="#F3F4F6")
    draw_centered_text(draw, "Triggers", 140, 440, 340, 560, f_title, fill="black")
    
    draw_arrow_with_label_offset(draw, 380, 220, 620, 220, "1", f_body, y_offset=-35)
    draw_arrow_with_label_offset(draw, 820, 220, 1100, 220, "N", f_body, y_offset=-35)
    
    draw_arrow_with_label_offset(draw, 240, 280, 240, 440, "1", f_body, x_offset=-30)
    draw_arrow_with_label_offset(draw, 240, 560, 240, 720, "N", f_body, x_offset=-30)
    
    draw_arrow_with_label_offset(draw, 380, 260, 620, 480, "1", f_body, x_offset=-40, y_offset=-30)
    draw_arrow_with_label_offset(draw, 820, 520, 1100, 740, "N", f_body, x_offset=40, y_offset=30)
    
    draw_arrow_with_label_offset(draw, 1270, 280, 1270, 440, "1", f_body, x_offset=-30)
    draw_arrow_with_label_offset(draw, 1270, 560, 1270, 720, "N", f_body, x_offset=-30)
    
    img.save("er.png")

# ==================== DOCX WORD GENERATOR STYLING HELPERS ====================

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{margin}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_left_border(cell, color_hex, size_pt):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(int(size_pt * 8)))
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    for b_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def set_table_borders(table, color_hex="D3D3D3"):
    tblPr = table._element.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_custom_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    
    run = h.runs[0]
    run.font.name = 'Calibri Light' if level == 1 else 'Calibri'
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(14)
    return h

def add_custom_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.65
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_list_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_bibliography_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def set_cell_text(cell, text, bold=False, size_pt=12):
    cell.text = text
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)

def add_callout(doc, text, title="NOTE", color_hex="000000", bg_hex="F3F4F6"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, bg_hex)
    set_cell_left_border(cell, color_hex, 3.0)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    
    run_title = p.add_run(f"[{title}] ")
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(12)
    run_title.font.color.rgb = RGBColor(0, 0, 0)
    
    run_text = p.add_run(text)
    run_text.font.name = 'Calibri'
    run_text.font.size = Pt(12)
    run_text.font.color.rgb = RGBColor(0, 0, 0)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)
    p_after.paragraph_format.line_spacing = 1.0

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F9FAFB")
    set_cell_left_border(cell, "000000", 2.5)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(6)
    p_after.paragraph_format.space_after = Pt(6)
    p_after.paragraph_format.line_spacing = 1.0

def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run1 = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run1._r.append(fldChar1)
        
        run2 = p.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run2._r.append(instrText)
        
        run3 = p.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar2)
        
        run4 = p.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run4._r.append(fldChar3)

# ==================== TABLE BUILDERS ====================

def add_schema_table(doc, fields):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], 'Field Name', bold=True)
    set_cell_text(hdr[1], 'Data Type', bold=True)
    set_cell_text(hdr[2], 'Constraint', bold=True)
    set_cell_text(hdr[3], 'Description', bold=True)
    for cell in hdr:
        set_cell_shading(cell, "F3F4F6")
        set_cell_margins(cell, 100, 100, 120, 120)
        
    for field in fields:
        row = table.add_row().cells
        for idx, text in enumerate(field):
            set_cell_text(row[idx], text)
        for cell in row:
            set_cell_margins(cell, 80, 80, 100, 100)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

def add_testing_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], 'Test Scenario', bold=True)
    set_cell_text(hdr[1], 'Input Details / Steps', bold=True)
    set_cell_text(hdr[2], 'Expected Behavior', bold=True)
    set_cell_text(hdr[3], 'Result', bold=True)
    for cell in hdr:
        set_cell_shading(cell, "F3F4F6")
        set_cell_margins(cell, 100, 100, 120, 120)
        
        test_cases = [
        ('TC-01: User Reg Pending State', 'PHO User signs up; status initialized to pending', 'Block access, show authorization warning', 'PASSED'),
        ('TC-02: User Reg Admin Approval', 'Admin clicks Approve on pending onboarding user', 'User state changes to approved; login allowed', 'PASSED'),
        ('TC-03: User Reg Rejection', 'Admin clicks Reject on pending account', 'User record updated to rejected; access blocked', 'PASSED'),
        ('TC-04: 2FA TOTP Key Generation', 'Approved User logs in first time; secret created', 'Show Base32 string and dynamic QR code screen', 'PASSED'),
        ('TC-05: 2FA Login Verification Pass', 'Enter Username + Password + valid 6-digit TOTP', 'Authenticate user and redirect to dashboard console', 'PASSED'),
        ('TC-06: 2FA Login Verification Fail', 'Enter Username + Password + invalid 6-digit TOTP', 'Block login, throw error warning in credentials screen', 'PASSED'),
        ('TC-07: Password Encryption Validation', 'Register User account; check DB entry field values', 'Verify password field contains 60-char bcrypt hash', 'PASSED'),
        ('TC-08: Automatic Inactivity Logout', 'Leave user session idle for 160 seconds on terminal', 'Trigger auto-logout alert; redirect to login screen', 'PASSED'),
        ('TC-09: Role Dashboard restrictions', 'Log in as Ship Agent; attempt to load /admin page', 'Redirect back to Agent dashboard; block access', 'PASSED'),
        ('TC-10: Vessel Registry Unique IMO Check', 'Insert vessel metadata; IMO: 9012345 (already exists)', 'Block saving, display duplicate registry error warning', 'PASSED'),
        ('TC-11: Vessel Registry Numeric Check', 'Insert GRT: "abcd" in vessel tonnage input fields', 'Block submission, trigger regex type validations', 'PASSED'),
        ('TC-12: Vessel Registry Length Check', 'Insert IMO: 12345 (5 digits) in registration form', 'Block registration; show 7-digit IMO limit alert', 'PASSED'),
        ('TC-13: Voyage Filing Blank Captain', 'File voyage logs; leave Captain name field blank', 'Block submission, highlight required captain inputs', 'PASSED'),
        ('TC-14: Voyage Filing Date Check', 'File arrival details; set ETD date before ETA date', 'Block voyage entry, display date logic constraint error', 'PASSED'),
        ('TC-15: Voyage Filing Cargo validation', 'File voyage details; set cargo category to general', 'Allow submission; save cargo classification successfully', 'PASSED'),
        ('TC-16: Stepper health approval step', 'Voyage state: PHO pending; Health officer clicks approve', 'Voyage state updates to Customs Pending; log audit', 'PASSED'),
        ('TC-17: Stepper customs audit step', 'Voyage state: Customs pending; Customs officer clicks approve', 'Voyage state updates to Traffic Pending; log audit', 'PASSED'),
        ('TC-18: Stepper traffic berthing step', 'Voyage state: Traffic pending; Traffic officer clicks approve', 'Voyage global status updates to Cleared; log audit', 'PASSED'),
        ('TC-19: Stepper clearance rejection', 'Voyage state: PHO pending; Health officer clicks reject', 'Voyage global status updates directly to Rejected', 'PASSED'),
        ('TC-20: Bilingual PDF Certificate Health', 'Click download health clearance certificate on cleared voyage', 'Generate PDF with green theme, bilingual headers and QR', 'PASSED'),
        ('TC-21: Bilingual PDF Certificate Port', 'Click download port clearance certificate on cleared voyage', 'Generate PDF with blue theme, CBIC emblems and QR', 'PASSED'),
        ('TC-22: Public QR Verification route', 'Scan QR on PDF with mobile; navigate to verify route', 'Fetch voyage status details without login; show valid', 'PASSED'),
        ('TC-23: Security Audit Trail Logging', 'Execute voyage clearance approval step in officer panel', 'Verify permanent log document created in audit collection', 'PASSED'),
        ('TC-24: Offline Fallback Connection drop', 'Turn off local network; execute vessel registry CRUD', 'Axios mock adapter intercepts; save data to localstorage', 'PASSED'),
        ('TC-25: Offline Fallback Connection sync', 'Reconnect network; trigger sync database handler console', 'Sync localstorage logs to MongoDB Atlas; clear local DB', 'PASSED'),
        ('TC-26: Language Toggle UI update', 'Click Hindi button on top bar; check page labels', 'Verify static labels update to Hindi characters successfully', 'PASSED'),
        ('TC-27: Accessibility Font Scaling Large', 'Click A+ button on top accessibility menu bar', 'Verify HTML root font increases by 2px with clean layout', 'PASSED'),
        ('TC-28: Accessibility Font Scaling Small', 'Click A- button on top accessibility menu bar', 'Verify HTML root font decreases by 2px without content clip', 'PASSED'),
        ('TC-29: Accessibility Theme Toggle Dark', 'Click theme toggle button; check body CSS data attribute', 'Verify theme sets to data-theme="dark" showing dark styles', 'PASSED'),
        ('TC-30: Sagar Setu Bot Panel Toggle', 'Click floating chatbot helper icon in dashboard console', 'Verify chatbot modal opens and collapses cleanly on clicks', 'PASSED'),
        ('TC-31: Sagar Setu Bot Quick Option Click', 'Click "How do I register a new vessel?" quick button', 'Verify chatbot prints reply detailing Vessel Registry steps', 'PASSED'),
        ('TC-32: Sagar Setu Bot Text Search query', 'Input "2fa authenticator" in message text input field', 'Verify chatbot matches keyword and answers with TOTP instructions', 'PASSED'),
        ('TC-33: Sagar Setu Bot Hindi Translation', 'Change page language to Hindi; check bot greeting message', 'Verify bot welcome message translates to Hindi characters', 'PASSED'),
        ('TC-34: SMS Gateway notification layout', 'Submit PHO clearance approval on voyage application', 'Verify floating SMS alert appears in corner with gateway header', 'PASSED'),
        ('TC-35: SMS Gateway alert auto dismiss', 'Wait 8 seconds after SMS gateway notification banner displays', 'Verify alert element is automatically removed from DOM', 'PASSED'),
        ('TC-36: Weather safety check wave limits', 'Mock NMPA wave height reading to 3.5m on wind panel', 'Verify weather widget shows red warning indicator log', 'PASSED'),
        ('TC-37: Weather safety check wind speed', 'Mock NMPA wind speed reading to 28 knots on panel', 'Verify weather widget displays high wind caution alert', 'PASSED'),
        ('TC-38: Carbon Estimator GRT computation', 'Change vessel Gross Tonnage value from 40k to 80k', 'Verify estimated carbon emission forecast updates dynamically', 'PASSED'),
        ('TC-39: Carbon Estimator cargo fuels check', 'Change voyage cargo category from Ballast to LNG cargo', 'Verify emissions estimator adjusts fuel factors accordingly', 'PASSED'),
        ('TC-40: Audit Logs Console Role Filter', 'Admin filters system audits logs by "Ship Agent" role', 'Verify log rows update to show actions from Agent accounts', 'PASSED'),
        ('TC-41: Audit Logs Console CSV Export', 'Click Export to CSV button on security logs panel', 'Verify browser creates download file containing spreadsheet logs', 'PASSED'),
        ('TC-42: Public Verify route invalid ID', 'Enter verify route URL with non-existent voyage ID', 'Verify page displays voyage not found warning modal', 'PASSED'),
        ('TC-43: Public Verify route valid clearance', 'Enter verify route URL with fully cleared voyage ID', 'Verify page displays verification status "VALID CLEARANCE"', 'PASSED'),
        ('TC-44: Session security logout alert', 'Wait 160 seconds without mouse movements or key taps', 'Verify modal alert states user has been logged out for safety', 'PASSED'),
        ('TC-45: Database connection drop banner check', 'Disconnect MongoDB cloud database; trigger API checks', 'Verify red alert banner warns offline local storage mock is active', 'PASSED'),
        ('TC-46: Sign Up Boundary Username Short', 'Sign up with username "ab" (2 chars); submit form', 'Verify client blocks validation; displays min length alert', 'PASSED'),
        ('TC-47: Sign Up Boundary Email Format', 'Sign up with email "agent-nmpa@gov"; missing domain suffix', 'Verify system rejects; triggers standard email validation pattern', 'PASSED'),
        ('TC-48: Sign Up Boundary Password Blank', 'Sign up with empty password value field', 'Verify submit button stays disabled until input field is satisfied', 'PASSED'),
        ('TC-49: Admin Panel Approve Approved Account', 'Admin attempts to approve an already approved user', 'Verify system ignores action; retains approved state without error', 'PASSED'),
        ('TC-50: Admin Panel User Search Filter', 'Admin types "Hel" in the search users console filter', 'Verify list updates to show only the health department account', 'PASSED'),
        ('TC-51: 2FA Secret Key Regenerate', 'Admin resets TOTP key for approved Ship Agent', 'Verify secret key changes to new Base32; 2FA prompt triggers', 'PASSED'),
        ('TC-52: 2FA TOTP Code Expired Check', 'Enter correct username, password and a 2-minute old TOTP code', 'Verify auth backend rejects code; returns 401 unauthorized log', 'PASSED'),
        ('TC-53: 2FA TOTP Non-Numeric characters', 'Enter credentials; type "abcdef" in 2FA TOTP input field', 'Verify inputs are restricted to numbers; block submit buttons', 'PASSED'),
        ('TC-54: Session Token Validation Check', 'API request sent with invalid JWT signature in header', 'Verify backend interceptor returns 403 authorization failed', 'PASSED'),
        ('TC-55: Session Token Expiration Check', 'API request sent with expired JWT session token in header', 'Verify backend responds 401 session expired; client redirects', 'PASSED'),
        ('TC-56: Vessel Registry Boundary IMO Format', 'Insert IMO: "912345A" containing non-numeric values', 'Verify regex checks block input; display IMO digit error', 'PASSED'),
        ('TC-57: Vessel Registry Boundary IMO Zero', 'Insert IMO: "0000000" in registration form', 'Verify database validation rejects IMO; triggers alert message', 'PASSED'),
        ('TC-58: Vessel Registry Boundary GRT Negative', 'Insert GRT: -120 in vessel tonnage input fields', 'Verify form validation blocks submission; requires value > 0', 'PASSED'),
        ('TC-59: Vessel Registry Boundary NRT Greater', 'Insert NRT greater than GRT tonnage details', 'Verify form blocks save; displays NRT cannot exceed GRT error', 'PASSED'),
        ('TC-60: Vessel Registry Boundary Owner Empty', 'Leave owner details blank; fill all other registry fields', 'Verify save button remains locked; highlights missing input', 'PASSED'),
        ('TC-61: Voyage Filing Boundary Port Origin', 'File voyage details with blank last port of call details', 'Verify submission blocked; highlights required origin inputs', 'PASSED'),
        ('TC-62: Voyage Filing Boundary ETA Format', 'File voyage with ETA set to invalid date formatting string', 'Verify client parser rejects; resets ETA value to current date', 'PASSED'),
        ('TC-63: Voyage Filing Boundary ETA Past', 'File voyage with ETA date set 30 days in the past', 'Verify form warning cautions agent on past date arrivals', 'PASSED'),
        ('TC-64: Voyage Filing Tonnage Fetch', 'Select vessel in voyage dropdown list; verify tonnage mapping', 'Verify GRT and NRT tonnage inputs auto-fill from registry values', 'PASSED'),
        ('TC-65: Stepper Action Double Approval PHO', 'PHO clicks approve button twice in rapid succession', 'Verify system processes first; ignores second to prevent duplicate logs', 'PASSED'),
        ('TC-66: Stepper Action Out-of-Sequence Customs', 'Customs attempts approval API request while voyage is PHO pending', 'Verify backend blocks state change; returns 400 workflow sequence error', 'PASSED'),
        ('TC-67: Stepper Action Out-of-Sequence Traffic', 'Traffic attempts approval API request while customs is pending', 'Verify backend blocks state change; returns 400 workflow sequence error', 'PASSED'),
        ('TC-68: Stepper Action Rejection Notes check', 'Customs clicks reject without writing audit comments', 'Verify system blocks action; prompts customs officer to enter comment details', 'PASSED'),
        ('TC-69: Stepper Action Comment length check', 'PHO inputs a 1000-character comment in approval notes', 'Verify comments save successfully; database text field expands', 'PASSED'),
        ('TC-70: jsPDF Download PDF Offline check', 'Agent clicks download certificate in offline fallback mode', 'Verify jsPDF triggers successfully; compiles PDF locally in browser', 'PASSED'),
        ('TC-71: jsPDF Verify QR Code parameters check', 'Scan QR code; check URL variables in query string', 'Verify URL contains correct voyage ID parameter for authentication', 'PASSED'),
        ('TC-72: Public Verify route SQL Injection', 'Input query string containing SQL characters in verify path', 'Verify backend sanitizes inputs; throws invalid format error', 'PASSED'),
        ('TC-73: Audit Logs Search query match', 'Admin searches audit logs for action "Vessel Created"', 'Verify logs table filters rows; shows only vessel creations', 'PASSED'),
        ('TC-74: Audit Logs Search no results check', 'Admin searches logs for non-existent keyword "payment"', 'Verify logs table renders empty state; displays no records found', 'PASSED'),
        ('TC-75: Offline Sync Storage Quota limit', 'Fill browser LocalStorage to 5.1MB; attempt vessel save', 'Verify client intercepts storage quota error; notifies user', 'PASSED'),
        ('TC-76: Offline Sync Conflict Resolution Cloud wins', 'Sync voyage logs where cloud DB has more recent timestamp', 'Verify database retains cloud values; updates LocalStorage copy', 'PASSED'),
        ('TC-77: Offline Sync State Indicator check', 'Check sync network status indicator in offline mode', 'Verify indicator widget displays "Offline mode - Local cache active"', 'PASSED'),
        ('TC-78: Weather Widget API failure check', 'Simulate openweather API timeout error in dashboard', 'Verify weather widget handles exception; displays fallback port warnings', 'PASSED'),
        ('TC-79: emissions calculator GRT zero check', 'Mock voyage vessel GRT to zero; trigger emissions check', 'Verify emissions estimator handles zero value; outputs zero emissions', 'PASSED'),
        ('TC-80: Sagar Setu Chatbot Null Input check', 'Click send button in chatbot console with empty text field', 'Verify chatbot ignores action; does not append empty bubbles', 'PASSED')
    ]
    for tc in test_cases:
        row = table.add_row().cells
        set_cell_text(row[0], tc[0])
        set_cell_text(row[1], tc[1])
        set_cell_text(row[2], tc[2])
        set_cell_text(row[3], tc[3], bold=True)
        for cell in row:
            set_cell_margins(cell, 80, 80, 100, 100)
            
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

def add_abbreviations_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], 'Abbreviation / Acronym', bold=True)
    set_cell_text(hdr[1], 'Definition', bold=True)
    for cell in hdr:
        set_cell_shading(cell, "F3F4F6")
        set_cell_margins(cell, 100, 100, 120, 120)
        
    ab_data = [
        ('NMPA', 'New Mangalore Port Authority'),
        ('SRS', 'Software Requirements Specification'),
        ('MERN', 'MongoDB, Express.js, React, Node.js Stack'),
        ('PHO', 'Port Health Organisation / Port Health Officer'),
        ('CBIC', 'Central Board of Indirect Taxes and Customs'),
        ('TOTP', 'Time-Based One-Time Password (2FA)'),
        ('JWT', 'JSON Web Token'),
        ('IMO', 'International Maritime Organization'),
        ('GRT / NRT', 'Gross Registered Tonnage / Net Registered Tonnage'),
        ('ILH', 'Indian Light House Dues'),
        ('SPA', 'Single Page Application'),
        ('CFD / DFD', 'Context Flow Diagram / Data Flow Diagram'),
        ('ER Model', 'Entity-Relationship Model'),
        ('CRUD', 'Create, Read, Update, Delete')
    ]
    for ab, definition in ab_data:
        row = table.add_row().cells
        set_cell_text(row[0], ab)
        set_cell_text(row[1], definition)
        for cell in row:
            set_cell_margins(cell, 80, 80, 100, 100)
            
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)

# ==================== PARSER AND APPENDER ====================

def append_section_from_file(doc, file_path):
    if not os.path.exists(file_path):
        print(f"Warning: file {file_path} not found.")
        return
        
    with open(file_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    parts = []
    current_pos = 0
    while True:
        start_idx = content.find("[CODE_START", current_pos)
        if start_idx == -1:
            parts.append(("text", content[current_pos:]))
            break
        parts.append(("text", content[current_pos:start_idx]))
        end_idx = content.find("[CODE_END]", start_idx)
        if end_idx == -1:
            parts.append(("text", content[start_idx:]))
            break
        code_block = content[start_idx:end_idx + 10]
        lang_start = code_block.find(":") + 1
        lang_end = code_block.find("]")
        code_text = code_block[lang_end + 2 : -10].strip()
        parts.append(("code", code_text))
        current_pos = end_idx + 10
        
    is_first_h2 = True
    for p_type, val in parts:
        if p_type == "code":
            add_code_block(doc, val)
        else:
            lines = val.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    add_custom_heading(doc, line[2:], level=1)
                elif line.startswith("## "):
                    if any(x in line for x in ["6.1.", "6.2.", "6.4.", "6.6.", "6.10.", "6.11.", "6.12."]):
                        if not is_first_h2:
                            doc.add_page_break()
                    add_custom_heading(doc, line[3:], level=2)
                    is_first_h2 = False
                elif line.startswith("### "):
                    if "Module" in line or "5.3." in line:
                        doc.add_page_break()
                    add_custom_heading(doc, line[4:], level=3)
                elif line.startswith("#### "):
                    add_custom_heading(doc, line[5:], level=4)
                elif line.startswith("- ") or line.startswith("• ") or line.startswith("* ") or (len(line) > 2 and line[0].isdigit() and line[1] == "."):
                    add_list_item(doc, line)
                elif line.startswith("[IMAGE:"):
                    img_path = line[7:-1].strip()
                    if os.path.exists(img_path):
                        p_img = doc.add_paragraph()
                        p_img.paragraph_format.line_spacing = 1.0
                        p_img.paragraph_format.space_before = Pt(0)
                        p_img.paragraph_format.space_after = Pt(0)
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        run_img.add_picture(img_path, width=Inches(5.5))
                        
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(4)
                        p_cap.paragraph_format.space_after = Pt(12)
                        cap_text = f"Figure: {os.path.basename(img_path)}"
                        run_cap = p_cap.add_run(cap_text)
                        run_cap.italic = True
                        run_cap.font.size = Pt(9.5)
                        run_cap.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        print(f"Warning: Screenshot image {img_path} not found.")
                else:
                    add_custom_paragraph(doc, text=line)

# ==================== MAIN COMPILER PIPELINE ====================

def build_report():
    print("Generating diagrams...")
    generate_architecture_diagram()
    generate_cfd()
    generate_dfd_l0()
    generate_dfd_l1()
    generate_dfd_l2()
    generate_decomposition_diagram()
    generate_er_diagram()
    print("Diagrams generated.")

    doc = docx.Document()
    
    # Configure 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.15)
        section.bottom_margin = Inches(1.15)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # ==================== COVER PAGE ====================
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(60)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(10)
    run_title = p_title.add_run("ACADEMIC PROJECT REPORT\nON\nPORT DIGITAL CLEARANCE SYSTEM FOR NMPA MANGALURU")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 0, 0)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(100)
    run_sub = p_sub.add_run("National Maritime Single Window Portal for New Mangalore Port Authority")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0, 0, 0)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(150)
    run_meta = p_meta.add_run("Submitted in partial fulfillment of the requirements for the degree of\n")
    run_meta.font.name = 'Calibri'
    run_meta.font.size = Pt(12)
    run_meta.font.color.rgb = RGBColor(0, 0, 0)
    run_meta = p_meta.add_run("Master of Computer Applications (MCA)\n")
    run_meta.font.bold = True
    run_meta.font.size = Pt(12)
    run_meta.font.color.rgb = RGBColor(0, 0, 0)
    
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(50)
    run_foot = p_foot.add_run("DEPARTMENT OF COMPUTER APPLICATIONS\nUNIVERSITY OF MANGALURU\nSubmission Year: 2026")
    run_foot.font.name = 'Calibri'
    run_foot.font.bold = True
    run_foot.font.size = Pt(12)
    run_foot.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_custom_heading(doc, "TABLE OF CONTENTS", level=1)
    
    table_toc = doc.add_table(rows=1, cols=3)
    table_toc.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_toc)
    
    hdr_cells = table_toc.rows[0].cells
    set_cell_text(hdr_cells[0], 'Chapter No', bold=True)
    set_cell_text(hdr_cells[1], 'Description', bold=True)
    set_cell_text(hdr_cells[2], 'Page No', bold=True)
    for cell in hdr_cells:
        set_cell_shading(cell, "F3F4F6")
        set_cell_margins(cell, 120, 120, 150, 150)
        
    toc_data = [
        ("1.", "Synopsis\n  Introduction\n  1.1. Introduction of the System\n      1.1.1. Project Title\n      1.1.2. Category\n      1.1.3. Overview\n  1.2. Background\n      1.2.1. Introduction of the Company\n  1.3. Objectives of the System\n  1.4. Scope of the System\n  1.5. Structure of the System\n  1.6. System Architecture\n  1.7. End Users\n  1.8. Software/Hardware used for the development\n  1.9. Software/Hardware required for the implementation", "4 - 15"),
        ("2.", "SRS\n  2.1. Introduction\n  2.2. Overall Description\n      2.2.1. Product perspective\n      2.2.2. Product Functions\n      2.2.3. User characteristics\n      2.2.4. General constraints\n      2.2.5. Assumptions\n  2.3. Special Requirements\n  2.4. Functional requirements\n  2.5. Design Constraints\n  2.6. System Attributes\n  2.7. Other Requirements", "16 - 29"),
        ("3.", "System Design\n  3.1. Introduction\n  3.2. Assumptions and Constraints\n  3.3. Functional decomposition\n  3.4. Description of Programs\n      3.4.1. Context Flow Diagram (CFD)\n      3.4.2. Data Flow Diagrams (DFDs - Level 0, Level 1, Level 2)", "30 - 41"),
        ("4.", "Database Design\n  4.1. Introduction\n  4.2. Purpose and scope\n  4.3. Table Definitions\n  4.4. ER Diagram", "42 - 53"),
        ("5.", "Detailed Design\n  5.1. Introduction\n  5.2. Structure of the software package\n  5.3. Modular decomposition of the System\n      5.3.1. Module 1: Authentication and Session Security\n      5.3.2. Module 2: Vessel Registry\n      5.3.3. Module 3: Voyage Clearance Stepper\n      5.3.4. Module 4: Bilingual jsPDF Generation\n      5.3.5. Module 5: Offline LocalStorage Sync Adapter", "54 - 65"),
        ("6.", "User Interface\n  6.1. Login\n  6.2. Main Screen/Home Page\n  6.3. Menu\n  6.4. Data Store/Retrieval/Update\n  6.5. Validation\n  6.6. View\n  6.7. On Screen reports\n  6.8. Alerts\n  6.9. Error message\n  6.10. Admin Panel Console\n  6.11. Audit Logs Console", "66 - 75"),
        ("7.", "Testing\n  7.1. Introduction\n      7.1.1. Unit Testing\n      7.1.2. Integrate Testing\n      7.1.3. System Testing\n  7.2. Test Reports", "76 - 79"),
        ("-", "Conclusion", "80"),
        ("-", "Limitations", "81"),
        ("-", "Scope for enhancement", "81"),
        ("-", "Abbreviation and Acronyms", "82"),
        ("-", "Bibliography", "82")
    ]
    
    for c_no, desc, p_no in toc_data:
        row = table_toc.add_row().cells
        set_cell_text(row[0], c_no)
        set_cell_text(row[1], desc)
        set_cell_text(row[2], p_no)
        for cell in row:
            set_cell_margins(cell, 80, 80, 100, 100)
            
    doc.add_page_break()

    # ==================== CHAPTER 1: SYNOPSIS ====================
    print("Compiling Chapter 1...")
    add_custom_heading(doc, "Chapter 1: Synopsis", level=1)
    append_section_from_file(doc, "report_sections/chapter_1_synopsis.txt")
    
    # Insert architecture diagram after 1.6 paragraph
    print("Inserting architecture diagram...")
    p_img = doc.add_paragraph()
    p_img.paragraph_format.line_spacing = 1.0
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(0)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture("architecture.png", width=Inches(5.8))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run("Figure 1.1: System Architecture Diagram")
    run_cap.italic = True
    run_cap.font.size = Pt(9.5)
    run_cap.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_page_break()

    # ==================== CHAPTER 2: SRS ====================
    print("Compiling Chapter 2...")
    add_custom_heading(doc, "Chapter 2: SRS", level=1)
    append_section_from_file(doc, "report_sections/chapter_2_srs.txt")
    doc.add_page_break()

    # ==================== CHAPTER 3: SYSTEM DESIGN ====================
    print("Compiling Chapter 3...")
    add_custom_heading(doc, "Chapter 3: System Design", level=1)
    
    # We parse chapter 3 file line by line
    with open("report_sections/chapter_3_system_design.txt", "r", encoding="utf-8") as f:
        ch3_content = f.read()
        
    ch3_parts = ch3_content.split("## ")
    for part in ch3_parts:
        if not part.strip():
            continue
        lines = part.split("\n")
        header = lines[0].strip()
        body = "\n".join(lines[1:]).strip()
        
        add_custom_heading(doc, header, level=2)
        
        # Add paragraphs
        sub_parts = body.split("\n\n")
        for sp in sub_parts:
            sp = sp.strip()
            if not sp:
                continue
            if sp.startswith("### "):
                add_custom_heading(doc, sp[4:], level=3)
            elif sp.startswith("- ") or sp.startswith("• "):
                # Split sub-list items
                list_items = sp.split("\n")
                for li in list_items:
                    add_list_item(doc, li.strip())
            else:
                add_custom_paragraph(doc, sp)
                
        # Insert diagrams programmatically at appropriate places
        if "3.3. Functional decomposition" in header:
            print("Inserting decomposition diagram...")
            p_img = doc.add_paragraph()
            p_img.paragraph_format.line_spacing = 1.0
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture("decomp.png", width=Inches(5.8))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run("Figure 3.1: Functional Decomposition Diagram")
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            
        elif "3.4. Description of Programs" in header:
            # Context Flow Diagram
            print("Inserting CFD diagram...")
            p_img = doc.add_paragraph()
            p_img.paragraph_format.line_spacing = 1.0
            p_img.paragraph_format.space_before = Pt(0)
            p_img.paragraph_format.space_after = Pt(0)
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture("cfd.png", width=Inches(5.5))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(12)
            run_cap = p_cap.add_run("Figure 3.2: Context Flow Diagram (CFD)")
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            
            # DFD Level 0
            print("Inserting DFD Level 0 diagram...")
            p_img2 = doc.add_paragraph()
            p_img2.paragraph_format.line_spacing = 1.0
            p_img2.paragraph_format.space_before = Pt(0)
            p_img2.paragraph_format.space_after = Pt(0)
            p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img2 = p_img2.add_run()
            run_img2.add_picture("dfd_l0.png", width=Inches(5.5))
            
            p_cap2 = doc.add_paragraph()
            p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap2.paragraph_format.space_before = Pt(4)
            p_cap2.paragraph_format.space_after = Pt(12)
            run_cap2 = p_cap2.add_run("Figure 3.3: Data Flow Diagram (DFD) Level 0")
            run_cap2.italic = True
            run_cap2.font.size = Pt(9.5)
            
            # DFD Level 1
            print("Inserting DFD Level 1 diagram...")
            p_img3 = doc.add_paragraph()
            p_img3.paragraph_format.line_spacing = 1.0
            p_img3.paragraph_format.space_before = Pt(0)
            p_img3.paragraph_format.space_after = Pt(0)
            p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img3 = p_img3.add_run()
            run_img3.add_picture("dfd_l1.png", width=Inches(5.8))
            
            p_cap3 = doc.add_paragraph()
            p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap3.paragraph_format.space_before = Pt(4)
            p_cap3.paragraph_format.space_after = Pt(12)
            run_cap3 = p_cap3.add_run("Figure 3.4: Data Flow Diagram (DFD) Level 1")
            run_cap3.italic = True
            run_cap3.font.size = Pt(9.5)
            
            # DFD Level 2
            print("Inserting DFD Level 2 diagram...")
            p_img4 = doc.add_paragraph()
            p_img4.paragraph_format.line_spacing = 1.0
            p_img4.paragraph_format.space_before = Pt(0)
            p_img4.paragraph_format.space_after = Pt(0)
            p_img4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img4 = p_img4.add_run()
            run_img4.add_picture("dfd_l2.png", width=Inches(5.5))
            
            p_cap4 = doc.add_paragraph()
            p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap4.paragraph_format.space_before = Pt(4)
            p_cap4.paragraph_format.space_after = Pt(12)
            run_cap4 = p_cap4.add_run("Figure 3.5: Data Flow Diagram (DFD) Level 2 - Stepper State Transitions")
            run_cap4.italic = True
            run_cap4.font.size = Pt(9.5)
            
    doc.add_page_break()

    # ==================== CHAPTER 4: DATABASE DESIGN ====================
    print("Compiling Chapter 4...")
    add_custom_heading(doc, "Chapter 4: Database Design", level=1)
    
    # We parse file content
    with open("report_sections/chapter_4_database_design.txt", "r", encoding="utf-8") as f:
        ch4_content = f.read()
        
    parts = []
    current_pos = 0
    while True:
        start_idx = ch4_content.find("[CODE_START", current_pos)
        if start_idx == -1:
            parts.append(("text", ch4_content[current_pos:]))
            break
        parts.append(("text", ch4_content[current_pos:start_idx]))
        end_idx = ch4_content.find("[CODE_END]", start_idx)
        if end_idx == -1:
            parts.append(("text", ch4_content[start_idx:]))
            break
        code_block = ch4_content[start_idx:end_idx + 10]
        lang_start = code_block.find(":") + 1
        lang_end = code_block.find("]")
        code_text = code_block[lang_end + 2 : -10].strip()
        parts.append(("code", code_text))
        current_pos = end_idx + 10
        
    for p_type, val in parts:
        if p_type == "code":
            add_code_block(doc, val)
        else:
            lines = val.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("## "):
                    add_custom_heading(doc, line[3:], level=2)
                elif line.startswith("### "):
                    sect_title = line[4:]
                    add_custom_heading(doc, sect_title, level=3)
                    
                    # Insert database schema tables programmatically
                    if "User Collection" in sect_title:
                        add_schema_table(doc, [
                            ('username', 'String', 'Required, Unique', 'Account handle name'),
                            ('password', 'String', 'Required', 'Bcrypt hashed password'),
                            ('email', 'String', 'Required', 'User email address'),
                            ('role', 'String', 'Required', 'Administrator, Agent, PHO, Customs, Traffic'),
                            ('status', 'String', 'Default: pending', 'pending, approved, rejected'),
                            ('twoFactorSecret', 'String', 'Optional', 'Base32 secret key for TOTP app'),
                            ('is2FAEnabled', 'Boolean', 'Default: true', 'Two-factor enforcement flag')
                        ])
                    elif "Vessel Collection" in sect_title:
                        add_schema_table(doc, [
                            ('name', 'String', 'Required', 'Name of registered vessel'),
                            ('imoNumber', 'String', 'Required, Unique', 'Unique 7-digit IMO registry key'),
                            ('flagState', 'String', 'Required', 'Country of registration flag'),
                            ('vesselType', 'String', 'Required', 'Container, Bulk, Oil Tanker, LPG, LNG'),
                            ('ownerDetails', 'String', 'Required', 'Shipping line owner description'),
                            ('grt', 'Number', 'Required', 'Gross Registered Tonnage value'),
                            ('nrt', 'Number', 'Required', 'Net Registered Tonnage value'),
                            ('userId', 'String', 'Optional', 'Creator agent account ID reference')
                        ])
                    elif "Voyage (Journey) Collection" in sect_title:
                        add_schema_table(doc, [
                            ('vessel', 'Object', 'Required', 'Snapshot of vessel metadata document'),
                            ('lastPortOfCall', 'String', 'Required', 'Last origin port of voyage'),
                            ('eta', 'Date', 'Required', 'Estimated Time of Arrival'),
                            ('etd', 'Date', 'Required', 'Estimated Time of Departure'),
                            ('status', 'String', 'Default: In Progress', 'In Progress, Cleared, Rejected'),
                            ('clearances', 'Object', 'Default: Pending', 'Sub-fields: health, customs, traffic statuses'),
                            ('notes', 'Object', 'Default: Empty', 'Sub-fields: health, customs, traffic comments'),
                            ('captainName', 'String', 'Optional', 'Master of the vessel'),
                            ('cargoType', 'String', 'Default: BALLAST', 'BALLAST, CONTAINER, CRUDE, LPG, LNG'),
                            ('crewCount', 'Number', 'Default: 0', 'Crew count on board'),
                            ('passengerCount', 'Number', 'Default: 0', 'Passenger count on board'),
                            ('ilhReceiptNo', 'String', 'Optional', 'Lighthouse dues receipt number'),
                            ('ilhAmount', 'Number', 'Default: 0', 'Lighthouse dues amount paid in INR')
                        ])
                    elif "Audit Trail Collection" in sect_title:
                        add_schema_table(doc, [
                            ('timestamp', 'Date', 'Default: Date.now', 'Creation date and time of record'),
                            ('user', 'String', 'Required', 'Username executing the action'),
                            ('action', 'String', 'Required', 'Description of system event logged')
                        ])
                elif line.startswith("- ") or line.startswith("• "):
                    add_list_item(doc, line)
                else:
                    add_custom_paragraph(doc, line)
                    
    # Insert ER diagram after 4.4 section text
    print("Inserting ER diagram...")
    p_img = doc.add_paragraph()
    p_img.paragraph_format.line_spacing = 1.0
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after = Pt(0)
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture("er.png", width=Inches(5.8))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(4)
    run_cap = p_cap.add_run("Figure 4.1: Entity Relationship (ER) Diagram")
    run_cap.italic = True
    run_cap.font.size = Pt(9.5)
    
    doc.add_page_break()

    # ==================== CHAPTER 5: DETAILED DESIGN ====================
    print("Compiling Chapter 5...")
    add_custom_heading(doc, "Chapter 5: Detailed Design", level=1)
    
    # We parse file content
    with open("report_sections/chapter_5_detailed_design.txt", "r", encoding="utf-8") as f:
        ch5_content = f.read()
        
    parts = []
    current_pos = 0
    while True:
        start_idx = ch5_content.find("[CODE_START", current_pos)
        if start_idx == -1:
            parts.append(("text", ch5_content[current_pos:]))
            break
        parts.append(("text", ch5_content[current_pos:start_idx]))
        end_idx = ch5_content.find("[CODE_END]", start_idx)
        if end_idx == -1:
            parts.append(("text", ch5_content[start_idx:]))
            break
        code_block = ch5_content[start_idx:end_idx + 10]
        lang_start = code_block.find(":") + 1
        lang_end = code_block.find("]")
        code_text = code_block[lang_end + 2 : -10].strip()
        parts.append(("code", code_text))
        current_pos = end_idx + 10
        
    for p_type, val in parts:
        if p_type == "code":
            add_code_block(doc, val)
        else:
            lines = val.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("## "):
                    add_custom_heading(doc, line[3:], level=2)
                elif line.startswith("### "):
                    if "Module" in line or "5.3." in line:
                        doc.add_page_break()
                    add_custom_heading(doc, line[4:], level=3)
                elif line.startswith("#### "):
                    add_custom_heading(doc, line[5:], level=4)
                elif line.startswith("- ") or line.startswith("• ") or line.startswith("* ") or (len(line) > 2 and line[0].isdigit() and line[1] == "."):
                    add_list_item(doc, line)
                else:
                    add_custom_paragraph(doc, line)
                    
                # Insert package structure after 5.2 section heading
                if "5.2. Structure of the software package" in line:
                    pkg_structure = (
                        "mca_project_remote/\n"
                        "├── backend/\n"
                        "│   ├── server.js (Main Express app entry)\n"
                        "│   ├── seed.js (Pre-populated role accounts script)\n"
                        "│   ├── models/ (Mongoose collections schemas)\n"
                        "│   │   ├── User.js\n"
                        "│   │   ├── Vessel.js\n"
                        "│   │   ├── Journey.js\n"
                        "│   │   └── AuditTrail.js\n"
                        "│   └── routes/\n"
                        "│       └── (REST controller API handlers)\n"
                        "└── frontend/\n"
                        "    ├── index.html\n"
                        "    └── src/\n"
                        "        ├── main.jsx (React initialization)\n"
                        "        ├── App.jsx (Navigation Router, Layout, Translation maps)\n"
                        "        ├── App.css (Global theme and styling)\n"
                        "        ├── mockBackend.js (LocalStorage offline database mock adapter)\n"
                        "        ├── Login.jsx (Credentials console and 2FA prompt)\n"
                        "        ├── VesselRegistry.jsx (Technical registrations CRUD)\n"
                        "        ├── ClearanceWorkflow.jsx (Clearance stepper forms & lists)\n"
                        "        ├── Dashboard.jsx (SVG charts and weather indicators)\n"
                        "        ├── AdminPanel.jsx (User approval queue)\n"
                        "        └── VerifyCertificate.jsx (Public QR verification router)"
                    )
                    add_code_block(doc, pkg_structure)
                    
    doc.add_page_break()

    # ==================== CHAPTER 6: USER INTERFACE ====================
    print("Compiling Chapter 6...")
    add_custom_heading(doc, "Chapter 6: User Interface", level=1)
    append_section_from_file(doc, "report_sections/chapter_6_ui.txt")
    doc.add_page_break()

    # ==================== CHAPTER 7: TESTING ====================
    print("Compiling Chapter 7...")
    add_custom_heading(doc, "Chapter 7: Testing", level=1)
    append_section_from_file(doc, "report_sections/chapter_7_testing.txt")
    
    # Insert Massive testing table programmatically
    print("Inserting 80 test cases table...")
    add_testing_table(doc)
    doc.add_page_break()

    # ==================== ENDING SECTIONS ====================
    print("Compiling Chapter 8 (Conclusion & Ending)...")
    append_section_from_file(doc, "report_sections/chapter_8_ending.txt")
    
    # Insert Abbreviations Table
    add_custom_heading(doc, "Abbreviation and Acronyms", level=1)
    add_abbreviations_table(doc)
    
    # Insert Bibliography list
    add_custom_heading(doc, "Bibliography", level=1)
    add_bibliography_item(doc, "1. Ministry of Ports, Shipping and Waterways, Government of India. Guidelines for National Maritime Single Window Portal, 2023.")
    add_bibliography_item(doc, "2. World Health Organization (WHO). International Health Regulations (IHR 2005) - Second Edition.")
    add_bibliography_item(doc, "3. Ministry of Health & Family Welfare, Government of India. Indian Port Health Rules, 1955.")
    add_bibliography_item(doc, "4. Central Board of Indirect Taxes and Customs (CBIC). The Customs Act, 1962 (India).")
    add_bibliography_item(doc, "5. Elmasri, R., & Navathe, S. B. Fundamentals of Database Systems, 7th Edition, Pearson.")
    add_bibliography_item(doc, "6. jsPDF Library Community, Client-Side JavaScript PDF Compilation Reference, GitHub documentation.")

    # Insert page numbers in footer
    add_page_number_footer(doc)

    # Save to local workspace directory
    out_path = "Port Digital Clearance System for NMPA Mangaluru_Project_Report.docx"
    doc.save(out_path)
    print(f"Project report successfully written to {out_path}")
    
    # Also save to D: download path if requested & directory exists
    alt_dir = "d:\\Download\\Port_Management_Website"
    if os.path.exists(alt_dir):
        try:
            alt_path = os.path.join(alt_dir, "Port Digital Clearance System for NMPA Mangaluru_Project_Report.docx")
            doc.save(alt_path)
            print(f"Project report successfully written to alt path {alt_path}")
        except Exception as e:
            print(f"Warning: could not save to D: path due to: {e}")
            
    # Clean up temp images
    for f in ["architecture.png", "cfd.png", "dfd_l0.png", "dfd_l1.png", "dfd_l2.png", "decomp.png", "er.png"]:
        if os.path.exists(f):
            try:
                os.remove(f)
            except Exception:
                pass

if __name__ == "__main__":
    build_report()
